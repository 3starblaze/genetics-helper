#!/usr/bin/env python3

from docx import Document
from gene import (
    Genotype, GeneAllele, breed, MultiGenotype,
    XXGenotype, XYGenotype, gender_breed
)
from typing import Sequence
from gene_convenience import create_genotype as g
from itertools import *

def create_table(genotype_a: Genotype, genotype_b: Genotype):
    document = Document()
    table = document.add_table(rows = 3, cols = 3)

    def set_cell(row, col, val):
        table.rows[row].cells[col].text = val

    set_cell(0, 0, "♀ \ ♂")

    set_cell(0, 1, genotype_b.gene_allele_a.value)
    set_cell(0, 2, genotype_b.gene_allele_b.value)
    set_cell(1, 0, genotype_a.gene_allele_a.value)
    set_cell(2, 0, genotype_a.gene_allele_b.value)

    breeding_results = [str(g) for g in breed(genotype_a, genotype_b)]

    set_cell(1, 1, breeding_results[0])
    set_cell(1, 2, breeding_results[1])
    set_cell(2, 1, breeding_results[2])
    set_cell(2, 2, breeding_results[3])

    document.save('sample.docx')

def create_multi_gene_table(mg_a: MultiGenotype, mg_b: MultiGenotype):
    document = Document()
    square_size = len(mg_a.genotypes) * len(mg_b.genotypes)

    table = document.add_table(rows = square_size + 1, cols = square_size + 1)

    def set_cell(row, col, val):
        table.rows[row].cells[col].text = val

    multi_genotype_breeding_results = mg_a.breed(mg_b)

    genotype_variations_a_temp = list(product(*[str(g) for g in mg_a.genotypes]))
    genotype_variations_b_temp = list(product(*[str(g) for g in mg_b.genotypes]))

    genotype_variations_a = ["".join(x) for x in genotype_variations_a_temp]
    genotype_variations_b = ["".join(x) for x in genotype_variations_b_temp]

    set_cell(0, 0, "♀ \ ♂")

    for i, genotype in enumerate(genotype_variations_a):
         set_cell(i + 1, 0, genotype)

    for i, genotype in enumerate(genotype_variations_b):
        set_cell(0, i + 1, genotype)

    element_counter = 0
    for row in range(square_size):
        for column in range(square_size):
            set_cell(row + 1, column + 1, multi_genotype_breeding_results[element_counter])
            element_counter += 1

    document.save('mgene_sample.docx')

def create_gender_table(female: XXGenotype, male: XYGenotype):
    document = Document()
    table = document.add_table(rows = 3, cols = 3)

    def set_cell(row, col, normal_value, superscript_value=None):
        p = table.rows[row].cells[col].add_paragraph(normal_value)
        if superscript_value == None: return
        p.add_run(superscript_value).font.superscript = True


    def get_info_text(f, m):
        return [
            ("X", f.gene_allele_a.value),
            ("X", f.gene_allele_b.value),
            ("X", m.gene_allele.value),
            ("Y"),
        ]

    def advanced_set_cell(row, col, gender_genotype):
        p = table.rows[row].cells[col].add_paragraph()
        if isinstance(gender_genotype, XXGenotype):
            p.add_run("X")
            p.add_run(gender_genotype.gene_allele_a.value).font.superscript = True
            p.add_run("X").font.superscript = False
            p.add_run(gender_genotype.gene_allele_b.value).font.superscript = True
        if isinstance(gender_genotype, XYGenotype):
            p.add_run("X")
            p.add_run(gender_genotype.gene_allele.value).font.superscript = True
            p.add_run("Y").font.superscript = False

    title_text = get_info_text(female, male)

    set_cell(0, 0, "♀ \ ♂")

    set_cell(0, 1, *title_text[2])
    set_cell(0, 2, *title_text[3])
    set_cell(1, 0, *title_text[0])
    set_cell(2, 0, *title_text[1])

    breeding_results = gender_breed(female, male)

    advanced_set_cell(1, 1, breeding_results[0])
    advanced_set_cell(1, 2, breeding_results[1])
    advanced_set_cell(2, 1, breeding_results[2])
    advanced_set_cell(2, 2, breeding_results[3])

    document.save('sample.docx')
